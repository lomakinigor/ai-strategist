"""Generate pipeline-overview.docx from the structured pipeline description."""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Default font
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# Title
title = doc.add_heading('Как работает приложение', level=1)
title.alignment = WD_ALIGN_PARAGRAPH.LEFT

# Intro
doc.add_paragraph(
    'Приложение представляет собой автоматизированный аналитический '
    'конвейер, который преобразует короткую анкету о компании в развёрнутый '
    'стратегический отчёт. Работа разбита на шесть последовательных этапов.'
)

# Numbered stages — using built-in List Number style for proper numbering
stages = [
    ('Сбор первичных данных.',
     'Владелец заполняет анкету: название, отрасль, сайт, ссылки на каналы '
     'в соцсетях, конкуренты, цели. Сведения сохраняются в базе данных.'),
    ('Параллельный сбор информации из открытых источников.',
     'Одновременно работают несколько потоков: интеллектуальный поисковик '
     'собирает упоминания компании, её рынка и конкурентов в открытом '
     'интернете, а отдельные адаптеры снимают метрики сайта (скорость, '
     'видимость в поиске), Telegram-канала (подписчики, охват, частота '
     'публикаций) и сообщества ВКонтакте.'),
    ('Оценка достоверности.',
     'Каждому факту присваивается числовая оценка от 1 до 5: государственные '
     'реестры и официальные сайты — 4–5, корпоративные соцсети — 3, '
     'рекламные материалы и отзывы — 2, неверифицированные агрегаторы — 1.'),
    ('Классификация утверждений.',
     'На основе этой оценки каждый факт относится к одной из трёх категорий: '
     'подтверждённый факт, гипотеза или недостаточно данных. Утверждения с '
     'низкой оценкой автоматически исключаются из дальнейшего анализа.'),
    ('Ручная верификация.',
     'Пользователь просматривает список собранных фактов и при необходимости '
     'отключает те, что считает некорректными или нерелевантными.'),
    ('Формирование отчёта.',
     'Подтверждённые факты группируются по пяти направлениям (бизнес, рынок, '
     'аудитория, каналы, конкуренты), после чего языковая модель составляет '
     'шесть содержательных разделов с практическими рекомендациями и оценкой '
     'потенциала автоматизации.'),
]

for i, (bold_part, rest) in enumerate(stages, start=1):
    p = doc.add_paragraph(style='List Number')
    run_bold = p.add_run(bold_part + ' ')
    run_bold.bold = True
    p.add_run(rest)

# Tools heading
doc.add_heading('Используемые инструменты', level=2)

# Tools table
tools = [
    ('Слой', 'API', 'Назначение', 'Тариф'),
    ('Research (web)', 'Perplexity Sonar',
     'Поиск фактов в реальном времени с цитатами', 'Платный'),
    ('Channels', 'VK API v5.199', 'Метрики ВКонтакте', 'Free tier'),
    ('Channels', 'Telegram t.me/s',
     'Парсер публичного превью канала', 'Бесплатно'),
    ('Site', 'Google PageSpeed v5',
     'Core Web Vitals, SEO, доступность', 'Free tier'),
    ('Strategy LLM', 'OpenRouter → DeepSeek-v4-flash',
     'Генерация разделов и синтез', 'Платный, дёшево'),
    ('DB', 'Postgres + pgvector на Neon',
     'Хранение фактов и отчётов', 'Free tier'),
]

table = doc.add_table(rows=len(tools), cols=4)
table.style = 'Light Grid Accent 1'

for row_idx, row in enumerate(tools):
    cells = table.rows[row_idx].cells
    for col_idx, value in enumerate(row):
        cells[col_idx].text = value
        # Bold header row + bold API column
        for para in cells[col_idx].paragraphs:
            for run in para.runs:
                if row_idx == 0:
                    run.bold = True
                elif col_idx == 1:  # API column
                    run.bold = True

out_path = 'docs/pipeline-overview.docx'
doc.save(out_path)
print(f'Saved: {out_path}')
